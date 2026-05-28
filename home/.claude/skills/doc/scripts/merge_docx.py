#!/usr/bin/env python3
"""合并三个修改版本到最终文档"""
import sys
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import shutil

def merge_documents(base_path, tech_path, layout_path, table_path, output_path):
    """
    合并三个修改：
    1. 从tech_path复制技术路线图(图2-1)
    2. 从layout_path复制图3-1
    3. 从table_path复制表4-1
    """

    # 加载所有文档
    base_doc = Document(base_path)
    tech_doc = Document(tech_path)
    layout_doc = Document(layout_path)
    table_doc = Document(table_path)

    print("=== 开始合并 ===\n")

    # 1. 替换技术路线图(图2-1) - 从tech_doc复制
    print("1. 查找并替换技术路线图(图2-1)...")
    replaced_tech = False
    for i, para in enumerate(base_doc.paragraphs):
        if '图2-1' in para.text or '图 2-1' in para.text:
            print(f"  找到图2-1位置: para {i}")
            # 找到tech_doc中的对应图片段落
            for tech_para in tech_doc.paragraphs:
                if tech_para._element.xpath('.//pic:pic'):
                    # 复制图片
                    for run in tech_para.runs:
                        if run._element.xpath('.//pic:pic'):
                            # 清空base_doc对应段落并复制
                            target_para = base_doc.paragraphs[i+1] if i+1 < len(base_doc.paragraphs) else para
                            target_para.clear()
                            for elem in run._element:
                                target_para._element.append(elem)
                            replaced_tech = True
                            print("  ✓ 技术路线图已替换")
                            break
                if replaced_tech:
                    break
            break

    # 2. 替换图3-1 - 从layout_doc复制
    print("\n2. 查找并替换图3-1...")
    replaced_layout = False
    for i, para in enumerate(base_doc.paragraphs):
        if '图3-1' in para.text or '图 3-1' in para.text:
            print(f"  找到图3-1位置: para {i}")
            # 找到layout_doc中的对应图片
            for layout_para in layout_doc.paragraphs:
                if layout_para._element.xpath('.//pic:pic'):
                    for run in layout_para.runs:
                        if run._element.xpath('.//pic:pic'):
                            target_para = base_doc.paragraphs[i+1] if i+1 < len(base_doc.paragraphs) else para
                            target_para.clear()
                            for elem in run._element:
                                target_para._element.append(elem)
                            replaced_layout = True
                            print("  ✓ 图3-1已替换")
                            break
                if replaced_layout:
                    break
            break

    # 3. 替换表4-1 - 从table_doc复制
    print("\n3. 查找并替换表4-1...")
    replaced_table = False

    # 在base_doc中找表4-1（查找"产品编号"表头且有年出入库量列）
    for i, element in enumerate(base_doc.element.body):
        if isinstance(element, CT_Tbl):
            table = Table(element, base_doc)
            if table.rows and len(table.rows) > 1:
                first_cell = table.rows[0].cells[0].text.strip()
                # 检查是否是产品年出入库量表
                if '产品编号' in first_cell and len(table.columns) >= 5:
                    header_row = [c.text.strip() for c in table.rows[0].cells]
                    if '年出入库量' in ''.join(header_row):
                        print(f"  找到表4-1位置: element {i}")

                        # 在table_doc中找对应表格
                        for j, table_elem in enumerate(table_doc.element.body):
                            if isinstance(table_elem, CT_Tbl):
                                new_table = Table(table_elem, table_doc)
                                if new_table.rows and len(new_table.rows) > 1:
                                    new_first = new_table.rows[0].cells[0].text.strip()
                                    if '产品编号' in new_first and len(new_table.columns) >= 5:
                                        new_header = [c.text.strip() for c in new_table.rows[0].cells]
                                        if '年出入库量' in ''.join(new_header):
                                            # 替换表格
                                            base_doc.element.body.remove(element)
                                            base_doc.element.body.insert(i, table_elem)
                                            replaced_table = True
                                            print("  ✓ 表4-1已替换")
                                            break
                        break

    # 保存
    base_doc.save(output_path)
    print(f"\n=== 合并完成 ===")
    print(f"输出文件: {output_path}")
    print(f"\n修改统计:")
    print(f"  技术路线图: {'✓' if replaced_tech else '✗'}")
    print(f"  图3-1: {'✓' if replaced_layout else '✗'}")
    print(f"  表4-1: {'✓' if replaced_table else '✗'}")

if __name__ == "__main__":
    base = "/Users/zhishixuebao/Desktop/外包项目/论文2/SLP 论文修改600/Matlab程序打包/杨松-基于SLP方法的k公司仓库布局优化研究-修改版(2).docx"
    tech = "/Users/zhishixuebao/Desktop/外包项目/论文2/SLP 论文修改600/Matlab程序打包/杨松-基于SLP方法的k公司仓库布局优化研究-修改版(3).docx"
    layout = "/Users/zhishixuebao/Desktop/外包项目/论文2/SLP 论文修改600/Matlab程序打包/杨松-基于SLP方法的k公司仓库布局优化研究-修改版(2)_图片优化.docx"
    table = "/Users/zhishixuebao/Desktop/外包项目/论文2/SLP 论文修改600/Matlab程序打包/杨松-基于SLP方法的k公司仓库布局优化研究-修改版(2)-已修改表4-1.docx"
    output = "/Users/zhishixuebao/Desktop/外包项目/论文2/SLP 论文修改600/Matlab程序打包/杨松-基于SLP方法的k公司仓库布局优化研究-最终版.docx"

    merge_documents(base, tech, layout, table, output)
