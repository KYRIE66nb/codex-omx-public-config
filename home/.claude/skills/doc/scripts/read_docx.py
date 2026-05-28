#!/usr/bin/env python3
"""Read and extract content from DOCX file."""
import sys
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def read_docx(path):
    doc = Document(path)

    print(f"=== Document: {path} ===\n")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total sections: {len(doc.sections)}\n")

    # Iterate through document body elements
    for i, element in enumerate(doc.element.body):
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else "Normal"
                print(f"[Para {i}] [{style}] {text[:100]}")

                # Check for images
                if para._element.xpath('.//pic:pic'):
                    print(f"  -> Contains image(s)")

        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            print(f"\n[Table {i}] {len(table.rows)} rows x {len(table.columns)} cols")
            # Show first few rows
            for r_idx, row in enumerate(table.rows[:3]):
                cells_text = [cell.text.strip()[:30] for cell in row.cells]
                print(f"  Row {r_idx}: {cells_text}")
            if len(table.rows) > 3:
                print(f"  ... ({len(table.rows) - 3} more rows)")
            print()

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: read_docx.py <path_to_docx>")
        sys.exit(1)
    read_docx(sys.argv[1])
